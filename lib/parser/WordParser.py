import docx
from lib.parser.BaseParser import BaseParser

class WordParser(BaseParser):
    def read(self):
        file = docx.Document(self.file_path)
        for para in file.paragraphs:
            para = para.text
            idcard_result = self.idcard_search(para)
            if idcard_result:
                self.sensitive_dict['idcard'].append(idcard_result)

            phone_result = self.phone_search(para)
            if phone_result:
                self.sensitive_dict['phone'].append(phone_result)

            email_result = self.email_search(para)
            if email_result:
                self.sensitive_dict['email'].append(email_result)

            ### 新增内容开始 ###
            sensitive_words = self.sensitive_word_search(para)
            if sensitive_words:
                self.sensitive_dict['sensitive_words'].extend(sensitive_words)
            ### 新增内容结束 ###

        tables = file.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    para = cell.text
                    idcard_result = self.idcard_search(para)
                    if idcard_result:
                        self.sensitive_dict['idcard'].append(idcard_result)

                    phone_result = self.phone_search(para)
                    if phone_result:
                        self.sensitive_dict['phone'].append(phone_result)

                    email_result = self.email_search(para)
                    if email_result:
                        self.sensitive_dict['email'].append(email_result)

                    ### 新增内容开始 ###
                    sensitive_words = self.sensitive_word_search(para)
                    if sensitive_words:
                        self.sensitive_dict['sensitive_words'].extend(sensitive_words)
                    ### 新增内容结束 ###

        self.reduce_error_report()
        return self.sensitive_dict